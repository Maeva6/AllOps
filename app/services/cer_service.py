import json
import re
from datetime import datetime
import fitz
from app.services.ai_errors import safe_chat_completion


# ══════════════════════════════════════════════════════════════════════
# EXTRACTION DU PROSIT ALLER (.docx / .md / .txt / .pdf)
# ══════════════════════════════════════════════════════════════════════

def extraire_texte_fichier(fichier) -> str:
    """Extrait le texte brut d'un fichier uploadé."""
    name      = fichier.filename.lower()
    raw_bytes = fichier.read()

    if name.endswith('.pdf'):
        doc   = fitz.open(stream=raw_bytes, filetype="pdf")
        texte = "\n".join(page.get_text() for page in doc)
        doc.close()
        return texte

    if name.endswith('.docx'):
        try:
            from docx import Document as DocxDoc
            import io
            doc = DocxDoc(io.BytesIO(raw_bytes))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            return f"[Erreur lecture docx : {e}]"

    return raw_bytes.decode('utf-8', errors='ignore')


def extraire_prosit_aller(texte: str) -> dict:
    """
    Parse le texte brut d'un prosit aller et retourne un dict
    avec les champs du formulaire CER pré-remplis.
    """
    champs = {
        'titre_prosit':     '',
        'numero_prosit':    '',
        'mots_cles':        '',
        'contexte':         '',
        'besoins':          '',
        'contraintes':      '',
        'problematique':    '',
        'pistes_solutions': '',
        'plan_action':      '',
    }

    SECTIONS = [
        ('mots_cles',        [r'mots[- ]cl[eé]s?', r'keywords?']),
        ('contexte',         [r'analyse du contexte', r'contexte', r'mise en situation', r'situation']),
        ('besoins',          [r'besoins?', r'notions? [àa] [eé]tudier', r'notions? cl[eé]s?', r'analyse des besoins']),
        ('contraintes',      [r'contraintes?']),
        ('problematique',    [r'd[eé]finition de la probl[eé]matique', r'probl[eé]matique', r'question centrale']),
        ('pistes_solutions', [r'pistes?(?: de solutions?)?', r'hypoth[eè]ses?']),
        ('plan_action',      [r"plan d[' ]action", r'plan de travail', r'[eé]tapes?']),
    ]

    all_titles = []
    for _, patterns in SECTIONS:
        all_titles.extend(patterns)
    boundary = '|'.join(all_titles)

    # Numéro et titre
    m_titre = re.search(
        r'PROSIT\s+(?:ALLER\s+)?N[°o]?\s*(\d+)\s*[:\-–—]\s*(.+)',
        texte, re.IGNORECASE
    )
    if m_titre:
        champs['numero_prosit'] = f"Prosit {m_titre.group(1)}"
        titre_brut = m_titre.group(2).strip()
        titre_brut = re.sub(
            r'\s*(cahier|cer|étudiant|pilote|co-pilote|promotion|date).*$',
            '', titre_brut, flags=re.IGNORECASE
        ).strip()
        champs['titre_prosit'] = f"Prosit N°{m_titre.group(1)} — {titre_brut}"
    else:
        for line in texte.splitlines():
            if line.strip():
                champs['titre_prosit'] = line.strip()
                break

    # Récupérer etudiant / pilote / copilote depuis le texte
    for label, key in [
        (r'[EÉ]tudiant\s*:', 'etudiant'),
        (r'Pilote\s*:', 'pilote'),
        (r'Co-pilote\s*:', 'copilote'),
        (r'Promotion\s*:', 'promotion'),
    ]:
        m = re.search(label + r'\s*(.+)', texte, re.IGNORECASE)
        if m:
            champs[key] = m.group(1).strip()

    # Extraire chaque section académique
    for champ, patterns in SECTIONS:
        for pat in patterns:
            regex = (
                r'(?:^|\n)'
                r'(?:[IVX]+\.\s*)?'
                r'(?:\d+[\.\)]\s*)?'
                r'(?:' + pat + r')\s*[:\-–]?\s*\n'
                r'(.*?)'
                r'(?=\n\s*(?:[IVX]+\.\s*)?(?:\d+[\.\)]\s*)?(?:' + boundary + r')\s*[:\-–]?\s*\n|$)'
            )
            m = re.search(regex, texte, re.IGNORECASE | re.DOTALL)
            if m:
                contenu = m.group(1).strip()
                contenu = re.sub(r'\n{3,}', '\n\n', contenu)
                champs[champ] = contenu
                break

    # Post-traitement plan_action
    if champs['plan_action']:
        lignes = []
        for ligne in champs['plan_action'].splitlines():
            ligne = ligne.strip()
            if ligne:
                ligne = re.sub(r'^[\d]+[\.\)]\s*', '', ligne)
                ligne = re.sub(r'^[-•*]\s*', '', ligne)
                if ligne:
                    lignes.append(ligne)
        champs['plan_action'] = '\n'.join(lignes)

    return champs


# ══════════════════════════════════════════════════════════════════════
# GÉNÉRATION WORD NATIVE (python-docx)
# Mise en page conforme au PDF exemple
# ══════════════════════════════════════════════════════════════════════

def generer_word_cer(cer) -> bytes:
    """
    Génère un .docx professionnel conforme à la mise en page du PDF exemple.
    Retourne les bytes du fichier.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io

    ROUGE = RGBColor(180, 0, 0)
    GRIS  = RGBColor(80,  80, 80)
    BLEU  = RGBColor(0,   70, 127)

    doc = Document()

    # Marges
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # Style Normal
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # ── Helpers ───────────────────────────────────────────────────────

    def set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

    def add_heading_rouge(text, level=1):
        h = doc.add_heading('', level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = h.add_run(text)
        run.font.color.rgb = ROUGE
        run.font.bold      = True
        run.font.size      = Pt({1: 14, 2: 12}.get(level, 11))
        return h

    def add_separator():
        p   = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '12')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), 'B40000')
        pBdr.append(bot)
        pPr.append(pBdr)

    def add_problematique_box(text):
        p   = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'DCE8F8')
        pPr.append(shd)
        p.paragraph_format.left_indent   = Cm(0.5)
        p.paragraph_format.right_indent  = Cm(0.5)
        p.paragraph_format.space_before  = Pt(6)
        p.paragraph_format.space_after   = Pt(6)
        # Bordure gauche rouge
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'),   'single')
        left.set(qn('w:sz'),    '18')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), 'B40000')
        pBdr.append(left)
        pPr.append(pBdr)
        run = p.add_run(text)
        run.font.italic = True
        run.font.size   = Pt(11)

    def add_code_block(lines):
        for line in (lines if lines else ['']):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  'F0F0F0')
            pPr.append(shd)
            run = p.add_run(line if line else ' ')
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.left_indent  = Cm(0.5)

    def inline_md(paragraph, text):
        """Applique gras/italique/code inline dans un paragraphe."""
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r = paragraph.add_run(part[2:-2]); r.bold = True
            elif part.startswith('*') and part.endswith('*'):
                r = paragraph.add_run(part[1:-1]); r.italic = True
            elif part.startswith('`') and part.endswith('`'):
                r = paragraph.add_run(part[1:-1])
                r.font.name = 'Courier New'; r.font.size = Pt(9)
            else:
                paragraph.add_run(part)

    def parse_md(text):
        """Parse le Markdown et ajoute au document."""
        if not text:
            return
        in_code  = False
        code_buf = []
        tbl_buf  = []
        in_table = False

        def flush_table():
            rows = [l for l in tbl_buf
                    if not re.match(r'^\|[-\s|:]+\|$', l.strip())]
            if not rows:
                return
            cols = [c.strip() for c in rows[0].strip('|').split('|')]
            nb   = len(cols)
            tbl  = doc.add_table(rows=len(rows), cols=nb)
            tbl.style = 'Table Grid'
            for i, row_txt in enumerate(rows):
                cells = [c.strip() for c in row_txt.strip('|').split('|')]
                while len(cells) < nb: cells.append('')
                for j, ct in enumerate(cells[:nb]):
                    cell = tbl.rows[i].cells[j]
                    cell.text = re.sub(r'\*\*(.+?)\*\*', r'\1', ct)
                    if i == 0:
                        for r2 in cell.paragraphs[0].runs:
                            r2.bold = True; r2.font.color.rgb = RGBColor(255,255,255)
                        set_cell_bg(cell, 'B40000')
            doc.add_paragraph()

        for line in text.splitlines():
            s = line.strip()

            if s.startswith('```'):
                if in_table: flush_table(); tbl_buf = []; in_table = False
                if in_code:
                    add_code_block(code_buf); code_buf = []; in_code = False
                else:
                    in_code = True
                continue
            if in_code:
                code_buf.append(line); continue

            if s.startswith('|'):
                in_table = True; tbl_buf.append(s); continue
            elif in_table:
                flush_table(); tbl_buf = []; in_table = False

            if not s:
                doc.add_paragraph(); continue

            if s.startswith('#### '):
                add_heading_rouge(s[5:], level=4)
            elif s.startswith('### '):
                add_heading_rouge(s[4:], level=3)
            elif s.startswith('## '):
                add_heading_rouge(s[3:], level=2)
            elif s.startswith('# '):
                add_heading_rouge(s[2:], level=1)
            elif re.match(r'^[-*+] ', s):
                p = doc.add_paragraph(style='List Bullet')
                inline_md(p, s[2:])
            elif re.match(r'^\d+\. ', s):
                p = doc.add_paragraph(style='List Number')
                inline_md(p, re.sub(r'^\d+\.\s+', '', s))
            else:
                p = doc.add_paragraph()
                inline_md(p, s)

        if in_code and code_buf:
            add_code_block(code_buf)
        if in_table and tbl_buf:
            flush_table()

    plan = cer.get_plan_action()
    date_str = (cer.created_at.strftime('%d/%m/%Y')
                if cer.created_at else datetime.now().strftime('%d/%m/%Y'))

    # ══════════════════════════════════════════════════════════════════
    # PAGE DE GARDE
    # ══════════════════════════════════════════════════════════════════

    # Titre UCAC-ICAM centré rouge
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run("UCAC-ICAM")
    r_logo.bold = True; r_logo.font.size = Pt(20); r_logo.font.color.rgb = ROUGE

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("La Fabrique à Génie")
    r_sub.font.size = Pt(11); r_sub.font.color.rgb = GRIS

    doc.add_paragraph()

    # Trait rouge
    add_separator()
    doc.add_paragraph()

    # Titre CER centré rouge
    p_cer = doc.add_paragraph()
    p_cer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cer = p_cer.add_run("CAHIER D'ÉTUDE ET DE RECHERCHE")
    r_cer.bold = True; r_cer.font.size = Pt(20); r_cer.font.color.rgb = ROUGE

    doc.add_paragraph()

    # Titre prosit dans encadré rouge (fond rose clair)
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr_t = p_tit._p.get_or_add_pPr()
    shd_t = OxmlElement('w:shd')
    shd_t.set(qn('w:val'),   'clear')
    shd_t.set(qn('w:color'), 'auto')
    shd_t.set(qn('w:fill'),  'FBEAEA')
    pPr_t.append(shd_t)
    # Bordure rouge tout autour
    pBdr_t = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '12')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), 'B40000')
        pBdr_t.append(el)
    pPr_t.append(pBdr_t)
    p_tit.paragraph_format.space_before = Pt(8)
    p_tit.paragraph_format.space_after  = Pt(8)
    r_tit = p_tit.add_run(cer.titre_prosit or 'CER')
    r_tit.bold = True; r_tit.font.size = Pt(16); r_tit.font.color.rgb = ROUGE

    doc.add_paragraph()
    doc.add_paragraph()

    # Tableau infos étudiant
    info_tbl = doc.add_table(rows=0, cols=2)
    info_tbl.style = 'Table Grid'

    def add_info_row(label, value):
        if not value: return
        row = info_tbl.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        for r2 in row.cells[0].paragraphs[0].runs:
            r2.bold = True
        set_cell_bg(row.cells[0], 'F0F0F0')

    add_info_row("Étudiant",  cer.etudiant  or '')
    add_info_row("Pilote",    cer.pilote    or '')
    add_info_row("Co-pilote", cer.copilote  or '')
    add_info_row("Promotion", cer.promotion or '')
    add_info_row("Date",      date_str)

    doc.add_paragraph()
    add_separator()
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # SECTIONS
    # ══════════════════════════════════════════════════════════════════

    def section(titre, contenu_fn):
        add_heading_rouge(titre, level=1)
        add_separator()
        contenu_fn()
        doc.add_paragraph()

    # I. Contexte
    section("I. Analyse du contexte", lambda: parse_md(cer.contexte))

    # II. Besoins
    add_heading_rouge("II. Analyse des besoins", level=1)
    add_separator()
    add_heading_rouge("Besoins", level=2)
    for b in (cer.besoins or '').splitlines():
        b = b.lstrip('-•* \t').strip()
        if b:
            p = doc.add_paragraph(style='List Bullet')
            inline_md(p, b)
    add_heading_rouge("Contraintes", level=2)
    for c in (cer.contraintes or '').splitlines():
        c = c.lstrip('-•* \t').strip()
        if c:
            p = doc.add_paragraph(style='List Bullet')
            inline_md(p, c)
    add_heading_rouge("Problématique", level=2)
    if cer.problematique:
        add_problematique_box(cer.problematique)
    doc.add_paragraph()

    # III. Généralisation
    add_heading_rouge("III. Généralisation", level=1)
    add_separator()
    p_gen = doc.add_paragraph()
    p_gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_gen = p_gen.add_run(cer.titre_prosit or '')
    r_gen.bold = True; r_gen.font.size = Pt(13); r_gen.font.color.rgb = ROUGE
    doc.add_paragraph()

    # IV. Plan d'action
    add_heading_rouge("IV. Plan d'action", level=1)
    add_separator()
    for i, etape in enumerate(plan, 1):
        p = doc.add_paragraph(style='List Number')
        inline_md(p, etape)
    doc.add_paragraph()

    # V. Réalisation
    add_heading_rouge("V. Réalisation du plan d'action", level=1)
    add_separator()
    if cer.realisation:
        try:
            real = json.loads(cer.realisation)
            for i, etape in enumerate(plan):
                add_heading_rouge(f"V.{i+1} {etape}", level=2)
                parse_md(real.get(str(i), ''))
        except Exception:
            parse_md(cer.realisation)
    doc.add_paragraph()

    # VI. Validation
    section("VI. Validation des pistes de solutions",
            lambda: parse_md(cer.validation or ''))

    # VII. Conclusion
    section("VII. Conclusion et retours sur les objectifs",
            lambda: parse_md(cer.conclusion or ''))

    # VIII. Bilan
    section("VIII. Bilan critique du travail effectué",
            lambda: parse_md(cer.bilan or ''))

    # IX. Synthèse
    if cer.synthese:
        section("IX. Synthèse des résultats obtenus",
                lambda: parse_md(cer.synthese or ''))

    # X. Références
    add_heading_rouge("Références bibliographiques", level=1)
    add_separator()
    if cer.references:
        items = re.findall(r'\[(\d+)\]\s+(.+?)(?=\n\s*\[\d+\]|\Z)',
                           cer.references, re.DOTALL)
        if items:
            for num, content in items:
                p = doc.add_paragraph()
                p.add_run(f"[{num}] ").bold = True
                inline_md(p, content.strip().replace('\n', ' '))
        else:
            parse_md(cer.references)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════
# UTILITAIRES LATEX
# ══════════════════════════════════════════════════════════════════════

def _esc(s: str) -> str:
    """Échappe les caractères spéciaux LaTeX pour les métadonnées."""
    if not s:
        return ''
    replacements = [
        ('\\', '\\textbackslash{}'),
        ('&',  '\\&'),
        ('%',  '\\%'),
        ('_',  '\\_'),
        ('#',  '\\#'),
        ('{',  '\\{'),
        ('}',  '\\}'),
        ('$',  '\\$'),
        ('^',  '\\^{}'),
        ('~',  '\\~{}'),
    ]
    for old, new in replacements:
        s = s.replace(old, new)
    return s


def nettoyer_texte_latex(text: str) -> str:
    if not text:
        return ""
    text = text.replace('\u2019', "'").replace('\u2018', "`")
    text = text.replace('\u201c', "``").replace('\u201d', "''")
    text = text.replace('\u2013', "--").replace('\u2014', "---")
    text = text.replace('\uf0d8', '')
    text = re.sub(r'(\d+)(?<!\\)%', r'\1\\%', text)
    return text


def nettoyer_code_listing(code: str) -> str:
    return ''.join(c if ord(c) <= 0x00FF else '?' for c in code)


# ══════════════════════════════════════════════════════════════════════
# GÉNÉRATION DES SECTIONS VIA GROQ
# ══════════════════════════════════════════════════════════════════════

def generer_section_cer(
    section: str,
    titre_prosit: str,
    contexte: str,
    besoins: str,
    contraintes: str,
    problematique: str,
    plan_action: list,
    contenu_source: str = "",
    etape_index: int = None,
    etape_label: str = None
) -> str:
    plan_str = "\n".join([f"{i+1}. {p}" for i, p in enumerate(plan_action)])
    source_ctx = (
        f"\nDocuments de référence (corbeille / workshop) :\n---\n{contenu_source[:4000]}\n---\n"
        if contenu_source else ""
    )

    if section == "realisation_etape":
        prompt = f"""Tu es un expert académique rédigeant un CER pour l'UCAC-ICAM.

PROSIT : {titre_prosit}
Contexte : {contexte}
Besoins : {besoins}
Contraintes : {contraintes}
Problématique : {problematique}
Plan d'action : {plan_str}
{source_ctx}

Rédige le contenu détaillé de l'ÉTAPE {etape_index} : "{etape_label}"

RÈGLES :
1. Minimum 400 mots, style académique rigoureux
2. Sous-sections (###) si nécessaire
3. Inclure : définitions formelles, exemples liés au prosit, tableaux comparatifs si pertinent,
   pseudocode/code si algorithmique, formules mathématiques ($...$) si pertinent
4. Markdown : **gras**, `code`, tableaux |col|col|, listes
5. Génère UNIQUEMENT le contenu, sans en-tête général
"""

    elif section == "validation":
        prompt = f"""Tu es un expert académique rédigeant un CER pour l'UCAC-ICAM.

PROSIT : {titre_prosit} | Problématique : {problematique}
Plan : {plan_str}
{source_ctx}

Rédige "Validation des pistes de solutions" :
- 4-5 sous-sections, chacune = une piste sous forme de question
- Format : sous-titre "Piste X : [question] ?" puis réponse 150-200 mots
- Conclure chaque piste : "**Piste validée / partiellement validée / invalidée.**"
"""

    elif section == "conclusion":
        prompt = f"""CER UCAC-ICAM — PROSIT : {titre_prosit}
Problématique : {problematique} | Plan : {plan_str}

Rédige "Conclusion et retours sur les objectifs" :
1. Rappel de la problématique
2. Bilan par étape (Atteint / Partiellement / Non atteint) en liste à puces
3. Synthèse de la réponse à la problématique
Minimum 250 mots.
"""

    elif section == "bilan":
        prompt = f"""CER UCAC-ICAM — PROSIT : {titre_prosit} | Plan : {plan_str}

Rédige "Bilan critique du travail effectué" en liste à puces :
- Points forts
- Limites et difficultés
- Perspectives et approfondissements possibles
Minimum 200 mots, style réflexif.
"""

    elif section == "synthese":
        prompt = f"""CER UCAC-ICAM — PROSIT : {titre_prosit} | Plan : {plan_str}

Rédige "Synthèse des résultats obtenus" :
- Résumé par étape en bullet points
- Données chiffrées si disponibles
- Résultats clés en **gras**
Minimum 200 mots.
"""

    elif section == "references":
        prompt = f"""CER UCAC-ICAM — PROSIT : {titre_prosit} | Thèmes : {plan_str}

Génère une bibliographie académique de 6-10 références.
Format OBLIGATOIRE (une référence par ligne) :
[1] Auteur, *Titre*, Éditeur, Année.
[2] ...
Inclure : livres académiques, articles, URLs réels (MIT, OpenClassrooms, docs officiels).
NE génère QUE la bibliographie, pas d'autre texte.
"""

    else:
        return ""

    response = safe_chat_completion(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
        max_tokens=3000,
    )
    return response.choices[0].message.content


# ══════════════════════════════════════════════════════════════════════
# CONVERSION MARKDOWN → LATEX
# Produit du LaTeX conforme au style du PDF exemple (boîtes tcolorbox,
# puces rouges, numéros rouges gras, booktabs, lstlisting)
# ══════════════════════════════════════════════════════════════════════

def markdown_to_latex(md_text: str) -> str:
    if not md_text:
        return ''
    md_text   = nettoyer_texte_latex(md_text)
    lines     = md_text.split('\n')
    result    = []
    in_code   = False
    in_item   = False
    in_enum   = False
    in_table  = False
    table_buf = []

    def pi(text):
        """Process inline markdown."""
        text = re.sub(r'\*\*(.+?)\*\*', r'\\textbf{\1}', text)
        text = re.sub(r'\*(.+?)\*',     r'\\textit{\1}', text)
        text = re.sub(r'`([^`]+)`',     r'\\texttt{\1}', text)
        text = re.sub(r'\[(.+?)\]\((.+?)\)', r'\\href{\2}{\1}', text)
        return text

    def close_lists():
        nonlocal in_item, in_enum
        out = []
        if in_item:
            out.append('\\end{itemize}')
            in_item = False
        if in_enum:
            out.append('\\end{enumerate}')
            in_enum = False
        return out

    def flush_table():
        nonlocal table_buf, in_table
        if not table_buf:
            return []
        # Filtrer les lignes séparatrices
        rows = [r for r in table_buf if not re.match(r'^\|[-\s|:]+\|$', r.strip())]
        if not rows:
            table_buf = []; in_table = False
            return []
        cols    = [c.strip() for c in rows[0].strip('|').split('|')]
        nb_cols = len(cols)
        out = [
            r'\begin{center}',
            r'\begin{tabular}{' + 'l' * nb_cols + '}',
            r'\toprule',
        ]
        for i, row in enumerate(rows):
            cells = [c.strip() for c in row.strip('|').split('|')]
            while len(cells) < nb_cols:
                cells.append('')
            if i == 0:
                out.append(' & '.join(r'\textbf{' + pi(c) + '}' for c in cells[:nb_cols]) + r' \\')
                out.append(r'\midrule')
            else:
                out.append(' & '.join(pi(c) for c in cells[:nb_cols]) + r' \\')
        out += [r'\bottomrule', r'\end{tabular}', r'\end{center}', '']
        table_buf = []; in_table = False
        return out

    LANG_MAP = {
        'python':'Python','java':'Java','c':'C','cpp':'C++',
        'js':'JavaScript','bash':'bash','sql':'SQL','':'Python'
    }

    for line in lines:
        s = line.strip()

        # ── Blocs de code ─────────────────────────────────────────────
        if s.startswith('```'):
            result.extend(close_lists())
            if in_table: result.extend(flush_table())
            if in_code:
                result += [r'\end{lstlisting}', '']
                in_code = False
            else:
                lang = LANG_MAP.get(s[3:].strip().lower(), 'Python')
                result.append(r'\begin{lstlisting}[language=' + lang + ']')
                in_code = True
            continue

        if in_code:
            result.append(nettoyer_code_listing(line))
            continue

        # ── Tableaux ──────────────────────────────────────────────────
        if s.startswith('|'):
            result.extend(close_lists())
            in_table = True
            table_buf.append(s)
            continue
        elif in_table:
            result.extend(flush_table())

        # ── Ligne vide ────────────────────────────────────────────────
        if not s:
            result.extend(close_lists())
            result.append('')
            continue

        # ── Titres ────────────────────────────────────────────────────
        if s.startswith('#### '):
            result.extend(close_lists())
            result.append(r'\paragraph{' + pi(s[5:]) + '}')
            continue
        if s.startswith('### '):
            result.extend(close_lists())
            result.append(r'\subsubsection{' + pi(s[4:]) + '}')
            continue
        if s.startswith('## '):
            result.extend(close_lists())
            result.append(r'\subsection{' + pi(s[3:]) + '}')
            continue
        if s.startswith('# '):
            result.extend(close_lists())
            result.append(r'\section{' + pi(s[2:]) + '}')
            continue

        # ── Listes non ordonnées (puces rouges comme dans le PDF) ─────
        if re.match(r'^[-*+] ', s):
            if not in_item:
                result.extend(close_lists())
                result.append(
                    r'\begin{itemize}[leftmargin=1.5cm,'
                    r' label=\textcolor{maincolor}{$\bullet$}]'
                )
                in_item = True
            result.append(r'  \item ' + pi(s[2:]))
            continue

        # ── Listes ordonnées (numéros rouges gras comme dans le PDF) ──
        if re.match(r'^\d+\. ', s):
            if not in_enum:
                result.extend(close_lists())
                result.append(
                    r'\begin{enumerate}[leftmargin=1.5cm,'
                    r' label=\textcolor{maincolor}{\textbf{\arabic*.}}]'
                )
                in_enum = True
            result.append(r'  \item ' + pi(re.sub(r'^\d+\.\s+', '', s)))
            continue

        # ── Paragraphe normal ─────────────────────────────────────────
        result.extend(close_lists())
        result.append(pi(s))
        result.append('')

    result.extend(close_lists())
    if in_table: result.extend(flush_table())
    if in_code:  result.append(r'\end{lstlisting}')

    text = '\n'.join(result)
    for env in ['itemize', 'enumerate']:
        diff = text.count(f'\\begin{{{env}}}') - text.count(f'\\end{{{env}}}')
        if diff > 0:
            text += f'\n\\end{{{env}}}' * diff
    return text


# ══════════════════════════════════════════════════════════════════════
# PRÉAMBULE LATEX  (identique au .tex exemple fourni)
# ══════════════════════════════════════════════════════════════════════

PREAMBLE = r"""\documentclass[12pt,a4paper]{report}
\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
\usepackage{geometry}
\usepackage{graphicx}
\usepackage{enumitem}
\usepackage{xcolor}
\usepackage{fancyhdr}
\usepackage{mdframed}
\usepackage{array}
\usepackage{booktabs}
\usepackage{longtable}
\usepackage{tabularx}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{listings}
\usepackage{float}
\usepackage{caption}
\usepackage{tikz}
\usetikzlibrary{positioning,shapes,arrows,calc}
\usepackage{url}
\usepackage{ragged2e}
\usepackage[colorlinks=true, linkcolor=black, urlcolor=blue, citecolor=black]{hyperref}
\usepackage{parskip}
\usepackage{titlesec}
\usepackage{multirow}
\usepackage{tcolorbox}
\tcbuselibrary{theorems,skins,breakable}

\geometry{margin=2.5cm, headheight=15pt}

% ===== COULEURS =====
\definecolor{maincolor}{RGB}{180,0,0}
\definecolor{lightgray}{RGB}{240,240,240}
\definecolor{codeblue}{RGB}{0,0,180}
\definecolor{lightblue}{RGB}{220,235,248}
\definecolor{lightgreen}{RGB}{230,255,230}
\definecolor{lightyellow}{RGB}{255,255,220}
\definecolor{lightred}{RGB}{255,240,240}
\definecolor{sectiongray}{RGB}{80,80,80}
\definecolor{darkgreen}{RGB}{0,120,0}

% ===== BOITES TCOLORBOX =====
\tcbset{
  defbox/.style={colback=lightblue,colframe=maincolor,
    fonttitle=\bfseries,arc=4pt,boxrule=1.5pt,breakable},
  proofbox/.style={colback=lightgreen,colframe=darkgreen,
    fonttitle=\bfseries\color{darkgreen},arc=4pt,boxrule=1pt,breakable},
  propbox/.style={colback=lightyellow,colframe=orange!70!black,
    fonttitle=\bfseries,arc=4pt,boxrule=1pt,breakable},
  warnbox/.style={colback=lightred,colframe=maincolor,
    fonttitle=\bfseries\color{maincolor},arc=4pt,boxrule=1pt,breakable}
}

% ===== STYLE CODE =====
\lstset{
  basicstyle=\small\ttfamily,
  backgroundcolor=\color{lightgray},
  frame=single, breaklines=true, language=Python,
  keywordstyle=\color{codeblue}\bfseries,
  commentstyle=\color{gray}\itshape,
  showstringspaces=false,
  numbers=left, numberstyle=\tiny\color{gray},
  stepnumber=1, tabsize=2,
  literate=
    {é}{{\'{e}}}1 {è}{{\`{e}}}1 {ê}{{\^{e}}}1 {ë}{{\"{e}}}1
    {à}{{\`{a}}}1 {â}{{\^{a}}}1 {ù}{{\`{u}}}1 {û}{{\^{u}}}1
    {î}{{\^{i}}}1 {ï}{{\"{i}}}1 {ô}{{\^{o}}}1 {ç}{{\c{c}}}1
    {É}{{\'{E}}}1 {È}{{\`{E}}}1 {À}{{\`{A}}}1 {Î}{{\^{I}}}1
    {Ç}{{\c{C}}}1 {œ}{{\oe}}1 {Œ}{{\OE}}1 {°}{{$^{\circ}$}}1,
}

% ===== EN-TÊTES FANCYHDR =====
\urlstyle{same}
\sloppy
\makeatletter
\g@addto@macro{\UrlBreaks}{\UrlOrds}
\makeatother

\renewcommand\thesection{\Roman{section}}
\renewcommand\thesubsection{\thesection.\arabic{subsection}}
"""


# ══════════════════════════════════════════════════════════════════════
# GÉNÉRATION DU LATEX COMPLET
# ══════════════════════════════════════════════════════════════════════

def generer_latex_complet(cer) -> str:
    """
    Génère le document LaTeX complet conforme au PDF exemple :
    CER_MAWAMBA_TOWA_Maëva_X2028.pdf
    """
    from datetime import datetime

    plan = cer.get_plan_action()

    date_str = (
        cer.created_at.strftime('%d/%m/%Y')
        if cer.created_at else datetime.now().strftime('%d/%m/%Y')
    )

    def md2tex(text):
        return markdown_to_latex(text) if text else ''

    # ── Titre : décomposer "Prosit NX — Thème" sur 2 lignes ───────────
    titre_raw = cer.titre_prosit or 'CER'
    m = re.match(
        r'^(Prosit\s+N[°o]?\s*\d+\s*[—\-–]+\s*.+?)\s*[—\-–]+\s*(.+)$',
        titre_raw, re.IGNORECASE
    )
    if m:
        t1 = _esc(m.group(1).strip())
        t2 = _esc(m.group(2).strip())
        titre_tikz_lines = (
            f"      {{\\Huge \\textbf{{\\underline{{{t1}}}}}}} \\\\[0.4cm]\n"
            f"      {{\\Huge \\textbf{{\\underline{{{t2}}}}}}} \\\\[0.6cm]\n"
        )
        theme_court = _esc(m.group(2).strip()[:38])
    else:
        t1 = _esc(titre_raw)
        titre_tikz_lines = (
            f"      {{\\Huge \\textbf{{\\underline{{{t1}}}}}}} \\\\[0.6cm]\n"
        )
        theme_court = _esc(titre_raw[:38])

    # ── Infos page de garde ───────────────────────────────────────────
    etudiant_esc  = _esc(cer.etudiant  or '')
    pilote_esc    = _esc(cer.pilote    or '')
    copilote_esc  = _esc(cer.copilote  or '')
    promotion_esc = _esc(cer.promotion or '')

    pilote_line   = f"  {{\\Large Pilote :{pilote_esc} \\par}}\n"   if cer.pilote   else ""
    copilote_line = f"  {{\\Large Co-pilote : {copilote_esc} \\par}}\n" if cer.copilote else ""
    promo_line    = f"  {{\\large {promotion_esc} \\par}}\n"         if cer.promotion else ""

    # ── Besoins / Contraintes → liste LaTeX ───────────────────────────
    def to_itemize_rouge(text: str) -> str:
        """Convertit un texte (une ligne = un item) en itemize rouge."""
        items = [
            l.lstrip('-•* \t').strip()
            for l in (text or '').splitlines()
            if l.strip() and not l.strip().startswith('#')
        ]
        if not items:
            return ''
        lines = [
            r'\begin{itemize}[leftmargin=1.5cm,'
            r' label=\textcolor{maincolor}{$\bullet$}]'
        ]
        for it in items:
            lines.append(r'  \item ' + _esc(it))
        lines.append(r'\end{itemize}')
        return '\n'.join(lines)

    # ── Plan d'action ─────────────────────────────────────────────────
    plan_items = '\n'.join(
        f'  \\item {_esc(p)}' for p in plan
    )

    # ── Réalisation ───────────────────────────────────────────────────
    realisation_sections = ""
    if cer.realisation:
        try:
            real = json.loads(cer.realisation)
            for i, etape in enumerate(plan):
                contenu = real.get(str(i), '')
                realisation_sections += (
                    "\n% " + "─" * 60 + "\n"
                    f"\\subsection{{{_esc(etape)}}}\n"
                    "% " + "─" * 60 + "\n\n"
                    + md2tex(contenu) + "\n\n"
                )
        except Exception:
            realisation_sections = md2tex(cer.realisation)

    # ── Validation → sous-sections ────────────────────────────────────
    # Le md2tex gère déjà les ### comme \subsubsection
    validation_tex = md2tex(cer.validation or '')

    # ── Bibliographie → \begin{thebibliography} ───────────────────────
    def build_biblio(text: str) -> str:
        if not text:
            return (
                r'\begin{thebibliography}{99}' + '\n'
                r'\bibitem{cormen} T.~H. Cormen, C.~E. Leiserson, R.~L. Rivest'
                r' et C.~Stein, \textit{Introduction \`{a} l\'algorithmique},'
                r' 3\`{e}me \'ed., Dunod, 2010.' + '\n'
                r'\end{thebibliography}'
            )
        items = re.findall(
            r'\[(\d+)\]\s+(.+?)(?=\n\s*\[\d+\]|\Z)',
            text, re.DOTALL
        )
        if not items:
            # Fallback : liste à puces → thebibliography auto-numérotée
            lignes = [l.lstrip('-•* ').strip() for l in text.splitlines() if l.strip()]
            out = [r'\begin{thebibliography}{99}']
            for j, l in enumerate(lignes, 1):
                l = re.sub(r'\*(.+?)\*', r'\\textit{\1}', l)
                out.append(f'\\bibitem{{ref{j}}} {l}')
            out.append(r'\end{thebibliography}')
            return '\n'.join(out)
        out = [r'\begin{thebibliography}{99}', '']
        for num, content in items:
            content = content.strip().replace('\n', ' ')
            content = re.sub(r'\*(.+?)\*', r'\\textit{\1}', content)
            out.append(f'\\bibitem{{ref{num}}}\n{content}\n')
        out.append(r'\end{thebibliography}')
        return '\n'.join(out)

    biblio_tex = build_biblio(cer.references or '')

    # ══════════════════════════════════════════════════════════════════
    # ASSEMBLAGE DU DOCUMENT
    # ══════════════════════════════════════════════════════════════════

    doc = PREAMBLE

    # En-têtes fancyhdr (dépend du titre → après PREAMBLE)
    doc += f"""
\\pagestyle{{fancy}}
\\fancyhf{{}}
\\fancyhead[L]{{\\textcolor{{maincolor}}{{\\textbf{{UCAC-ICAM}}}}}}
\\fancyhead[R]{{\\textcolor{{maincolor}}{{\\textbf{{CER -- {theme_court}}}}}}}
\\fancyfoot[C]{{\\textcolor{{sectiongray}}{{\\thepage}}}}
\\renewcommand{{\\headrulewidth}}{{1.5pt}}
\\renewcommand{{\\headrule}}{{%
  \\hbox to\\headwidth{{\\color{{maincolor}}\\leaders\\hrule height \\headrulewidth\\hfill}}}}

\\begin{{document}}

% ============================================================
%                     PAGE DE GARDE
% ============================================================
\\begin{{titlepage}}
  \\pagestyle{{empty}}
  % Cadre rouge épais tout autour (identique au PDF exemple)
  \\begin{{tikzpicture}}[remember picture, overlay]
    \\draw[line width=24pt, color=maincolor]
      (current page.north west) rectangle (current page.south east);
  \\end{{tikzpicture}}
  \\centering
  \\vspace{{1cm}}

  % ── Logo ──────────────────────────────────────────────────────────
  % Utilise logo.jpg si présent dans le même dossier que le .tex
  % Sinon, bloc TikZ de remplacement (commenter/décommenter selon le cas)
  %\\includegraphics[width=5cm]{{logo.jpg}}
  \\begin{{tikzpicture}}
    \\fill[maincolor, rounded corners=6pt] (0,0) rectangle (5.5,2.2);
    \\fill[white, rounded corners=3pt] (0.15,0.15) rectangle (5.35,2.05);
    \\node[font=\\Large\\bfseries, text=maincolor] at (2.75,1.6) {{UCAC-ICAM}};
    \\node[font=\\footnotesize, text=sectiongray] at (2.75,1.0)
      {{Universit\\\'{{e}} Catholique d'Afrique Centrale}};
    \\node[font=\\footnotesize, text=sectiongray] at (2.75,0.55)
      {{Institut Catholique d'Arts et M\\\'{{e}}tiers}};
    \\draw[maincolor, line width=0.6pt] (0.9,1.3) -- (4.6,1.3);
  \\end{{tikzpicture}}

  \\vspace{{1.5cm}}

  % ── Titre dans encadré TikZ arrondi (identique au PDF exemple) ────
  \\begin{{tikzpicture}}
    \\node[draw=maincolor, line width=1.5pt, rounded corners=5pt,
          inner sep=16pt, align=center] {{
{titre_tikz_lines}      {{\\Huge \\textbf{{Cahier d\\'\\\'{{E}}tude}}}} \\\\[0.3cm]
      {{\\Huge \\textbf{{et de Recherche}}}}
    }};
  \\end{{tikzpicture}}

  \\vfill

  % ── Informations étudiant / encadrants ────────────────────────────
  {{\\Large \\textbf{{{etudiant_esc}}} \\par}}
  \\vspace{{0.4cm}}
  {{\\large {date_str} \\par}}
  \\vspace{{0.6cm}}
{pilote_line}{copilote_line}{promo_line}\\end{{titlepage}}

% ============================================================
%                  TABLE DES MATIÈRES
% ============================================================
{{
\\renewcommand{{\\baselinestretch}}{{2.0}}
\\tableofcontents
}}
\\newpage

% ============================================================
\\section{{Analyse du contexte}}
% ============================================================

{md2tex(cer.contexte)}

% ============================================================
\\section{{Analyse des besoins}}
% ============================================================

\\subsection*{{Besoins}}
{to_itemize_rouge(cer.besoins)}

\\subsection*{{Contraintes}}
{to_itemize_rouge(cer.contraintes)}

\\subsection*{{Probl\\\'{{e}}matique}}

\\begin{{mdframed}}[backgroundcolor=lightblue, linecolor=maincolor, linewidth=2pt,
  innerleftmargin=10pt, innerrightmargin=10pt,
  innertopmargin=8pt, innerbottommargin=8pt]
\\textit{{{_esc(cer.problematique or '')}}}
\\end{{mdframed}}

% ============================================================
\\section{{G\\\'{{e}}n\\\'{{e}}ralisation}}
% ============================================================

\\begin{{center}}
  \\large\\textbf{{\\textcolor{{maincolor}}{{{_esc(titre_raw)}}}}}
\\end{{center}}

Ce prosit s\\'inscrit dans un domaine fondamental de l\\'informatique et des
math\\\'{{e}}matiques appliqu\\\'{{e}}es. Les notions \\\'{{e}}tudi\\\'{{e}}es ont
une port\\\'{{e}}e th\\\'{{e}}orique et pratique d\\\'{{e}}passant le seul cadre de ce prosit.

% ============================================================
\\section{{Plan d\\'action}}
% ============================================================

\\begin{{enumerate}}[leftmargin=1.5cm, label=\\textcolor{{maincolor}}{{\\textbf{{\\arabic*.}}}}]
{plan_items}
\\end{{enumerate}}

% ============================================================
\\section{{R\\\'{{e}}alisation du plan d\\'action}}
% ============================================================

{realisation_sections}

% ============================================================
\\section{{Validation des pistes de solutions}}
% ============================================================

{validation_tex}

% ============================================================
\\section{{Conclusion et retours sur les objectifs}}
% ============================================================

{md2tex(cer.conclusion or '')}

% ============================================================
\\section{{Bilan critique du travail effectu\\\'{{e}}}}
% ============================================================

{md2tex(cer.bilan or '')}

% ============================================================
%                 BIBLIOGRAPHIE
% ============================================================

{biblio_tex}

\\end{{document}}
"""
    return doc